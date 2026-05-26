#!/usr/bin/env python3
"""Generate the Vietnamese user guide (.docx) for EduGate.

Run: python3 tools/gen-user-guide.py
Output: docs/HUONG_DAN_SU_DUNG_EduGate.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

PRIMARY = RGBColor(0x1D, 0x4E, 0xD8)  # blue-700-ish


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = PRIMARY
    return p


def h2(text):
    return doc.add_heading(text, level=2)


def para(text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def step(text):
    return doc.add_paragraph(text, style="List Number")


def note(text):
    p = doc.add_paragraph()
    r = p.add_run("Lưu ý: ")
    r.bold = True
    r.font.color.rgb = PRIMARY
    p.add_run(text)
    return p


# ---- Cover ----------------------------------------------------------------
title = doc.add_heading("HƯỚNG DẪN SỬ DỤNG", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("EduGate — Hệ thống tuyển sinh trực tuyến đa trường")
r.bold = True
r.font.size = Pt(14)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Phiên bản 1 (Module Tuyển sinh) · Cập nhật: 26/05/2026").italic = True
doc.add_paragraph()

# ---- 1. Giới thiệu --------------------------------------------------------
h1("1. Giới thiệu")
para(
    "EduGate là nền tảng tuyển sinh trực tuyến dùng chung cho nhiều trường. Mỗi "
    "trường là một không gian riêng (tenant) với giao diện, màu sắc và biểu mẫu "
    "tuyển sinh riêng. Hệ thống phục vụ hai nhóm người dùng:"
)
bullet("Phụ huynh / học sinh: đăng ký tuyển sinh, theo dõi và tra cứu hồ sơ, tải biên lai.")
bullet("Cán bộ nhà trường (quản trị): tiếp nhận, tìm kiếm, lọc, duyệt hồ sơ và xem thống kê.")

# ---- 2. Truy cập ----------------------------------------------------------
h1("2. Truy cập hệ thống")
para("Mỗi trường có một đường dẫn riêng theo mã trường:")
bullet("Trang tuyển sinh của trường: /t/<mã-trường>  (ví dụ: /t/nguyen-huy-tuong)")
bullet("Khu quản trị của trường: /t/<mã-trường>/admin")
para("Một số trường đang có trên hệ thống demo:")
bullet("Trường Tiểu học & THCS Cầu Vàng — mã: cva-edu")
bullet("Trường THPT Trần Đại Nghĩa — mã: tran-dai-nghia")
bullet("Trường THCS Nguyễn Huy Tưởng (Đông Anh) — mã: nguyen-huy-tuong")
bullet("Trường THCS Nguyễn Văn Huyên (Hoài Đức) — mã: nguyen-van-huyen")
bullet("Trường THCS Nguyễn Gia Thiều (Long Biên) — mã: nguyen-gia-thieu")
note(
    "Giao diện được tối ưu cho cả điện thoại và máy tính. Trên điện thoại, thanh "
    "thao tác (Tiếp tục / Quay lại) được ghim ở đáy màn hình cho dễ bấm."
)

# ---- 3. Dành cho phụ huynh ------------------------------------------------
h1("3. Dành cho phụ huynh")

h2("3.1. Đăng ký tuyển sinh")
para(
    "Vào trang trường rồi bấm \"Đăng ký ngay\" (hoặc mở /t/<mã-trường>/register). "
    "Quy trình gồm 4 bước, có thanh tiến trình hiển thị bước hiện tại. Thông tin "
    "được tự động lưu nháp trên thiết bị nên bạn có thể quay lại điền tiếp."
)
para("Bước 1 — Người khai:", bold=True)
step("Nhập Họ và tên, Email, Số điện thoại của người khai (phụ huynh/người giám hộ).")
step("Chọn Quan hệ với học sinh (Bố / Mẹ / Người giám hộ / Bản thân học sinh).")
step("Các trường có dấu sao (*) màu đỏ là bắt buộc. Điền đủ và đúng định dạng thì nút \"Tiếp tục\" mới sáng lên.")
para("Bước 2 — Hồ sơ:", bold=True)
step("Nhập Mã học sinh (do Bộ GD&ĐT cấp), rồi bấm \"Tra cứu\". Hệ thống tự hiển thị Họ tên, Ngày sinh, Giới tính của học sinh.")
step("Tải Hồ sơ học bạ lên: kéo-thả tệp vào vùng nét đứt, hoặc bấm để chọn tệp (chấp nhận PDF, JPG, PNG). Khi tải đúng định dạng sẽ có thông báo thành công.")
step("Có thể nhập \"Ghi chú thêm\" (không bắt buộc).")
para("Bước 3 — Xác thực email:", bold=True)
step("Hệ thống gửi mã OTP gồm 6 chữ số tới email người khai.")
step("Nhập mã vào 6 ô vuông, hoặc dán (paste) trực tiếp mã copy từ email — hệ thống tự chia vào các ô.")
step("Bấm \"Xác thực & nộp hồ sơ\". Nếu cần, dùng \"Gửi lại mã\".")
para("Bước 4 — Hoàn tất:", bold=True)
step("Màn hình hiển thị lời cảm ơn kèm tên phụ huynh và Mã hồ sơ.")
step("Hãy lưu lại Mã hồ sơ để tra cứu trạng thái về sau.")
step("Có thể bấm \"Tra cứu hồ sơ\" hoặc \"Tải biên lai\" ngay tại đây.")
note("Trong bản demo, mã OTP xác thực email là 123456 (hiển thị kèm ở thông báo).")

h2("3.2. Tra cứu hồ sơ")
step("Mở /t/<mã-trường>/track (hoặc bấm \"Tra cứu hồ sơ\" trên trang trường).")
step("Nhập Mã hồ sơ đã nhận khi đăng ký, bấm \"Tra cứu\".")
step("Xem trạng thái hiện tại và lịch sử xử lý hồ sơ theo thời gian.")

h2("3.3. Tải biên lai và hồ sơ (PDF)")
para(
    "Tại trang tra cứu hồ sơ, bạn có thể tải/in: Biên lai nộp hồ sơ và Hồ sơ tuyển "
    "sinh đầy đủ. Biên lai ghi rõ số thứ tự hồ sơ đã tiếp nhận, mã hồ sơ, thông "
    "tin người khai, trạng thái, và mang logo + màu sắc của trường."
)

# ---- 4. Dành cho cán bộ nhà trường ---------------------------------------
h1("4. Dành cho cán bộ nhà trường (quản trị)")

h2("4.1. Đăng nhập")
step("Mở /t/<mã-trường>/admin hoặc trang đăng nhập /login.")
step("Đăng nhập bằng email cán bộ được cấp (ví dụ: tuyensinh.nht@edu.vn cho trường Nguyễn Huy Tưởng).")
note("Bản demo chấp nhận đăng nhập với thông tin hợp lệ bất kỳ và không chặn truy cập thủ công.")

h2("4.2. Bảng điều khiển (Dashboard)")
para("Trang chủ khu quản trị tổng hợp nhanh tình hình tuyển sinh:")
bullet("Số hồ sơ nộp hôm nay, đang duyệt, cần bổ sung, đã duyệt.")
bullet("Biểu đồ phân bố hồ sơ theo trạng thái và lượng nộp theo ngày (14 ngày gần nhất).")
bullet("Phễu chuyển đổi và danh sách hồ sơ mới nhất.")

h2("4.3. Quản lý hồ sơ")
para("Vào mục \"Hồ sơ\" (/t/<mã-trường>/admin/applications):")
bullet("Tìm kiếm: theo mã hồ sơ, tên người khai hoặc tên học sinh.")
bullet("Lọc: theo trạng thái (chọn nhiều), theo khoảng ngày, theo khoảng điểm; có thẻ hiển thị các bộ lọc đang áp dụng.")
bullet("Phân trang: chọn số dòng mỗi trang và chuyển trang.")
bullet("Trên máy tính hiển thị dạng bảng; trên điện thoại hiển thị dạng thẻ cho dễ đọc.")
bullet("Bấm vào một hồ sơ để mở chi tiết: thông tin người khai, nội dung biểu mẫu, lịch sử trạng thái và các thao tác đổi trạng thái (duyệt, yêu cầu bổ sung, từ chối...).")

h2("4.4. Thông báo")
para(
    "Biểu tượng chuông ở thanh trên cùng hiển thị các thông báo liên quan tới việc "
    "tiếp nhận hồ sơ (hồ sơ mới, cần bổ sung tài liệu, sắp hết hạn...), kèm số "
    "thông báo chưa đọc. Có thể đánh dấu đã đọc từng mục hoặc tất cả."
)

h2("4.5. Tùy chỉnh thương hiệu & trang đích")
para(
    "Trong \"Cài đặt\", trường có thể chỉnh màu sắc thương hiệu, tên hiển thị và nội "
    "dung các khối trên trang tuyển sinh (hero, giới thiệu, quy trình, FAQ...) với "
    "bản xem trước trực tiếp — không cần lập trình hay triển khai lại."
)

# ---- 5. FAQ ---------------------------------------------------------------
h1("5. Câu hỏi thường gặp")
para("Quên mã hồ sơ?", bold=True)
para("Mã hồ sơ được gửi qua email đã xác thực khi nộp. Hãy tìm trong hộp thư của bạn.")
para("Có sửa được hồ sơ sau khi nộp?", bold=True)
para("Khi hồ sơ ở trạng thái \"Cần bổ sung\", phụ huynh có thể cập nhật và nộp lại theo hướng dẫn của nhà trường.")
para("Nút \"Tiếp tục\" không bấm được?", bold=True)
para("Kiểm tra lại các trường bắt buộc (dấu * đỏ) đã điền đủ và đúng định dạng (email, số điện thoại) chưa.")

# ---- 6. Ghi chú phiên bản -------------------------------------------------
h1("6. Ghi chú phiên bản")
bullet("Đây là bản pha 1, tập trung Module Tuyển sinh; dữ liệu hồ sơ trong khu quản trị hiện là dữ liệu mô phỏng phục vụ trình diễn.")
bullet("Mã OTP demo: 123456. Mã học sinh dạng 8–12 chữ số sẽ tự sinh thông tin mô phỏng.")
bullet("Các phân hệ khác (Nhân sự, Học phí, LMS, Khảo sát...) sẽ bổ sung theo nhu cầu.")

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.add_run("© EduGate — Tài liệu hướng dẫn sử dụng").italic = True

import os
os.makedirs("docs", exist_ok=True)
out = "docs/HUONG_DAN_SU_DUNG_EduGate.docx"
doc.save(out)
print("Saved:", out)
